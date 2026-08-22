#!/usr/bin/env python3
"""build_docx.py — 6 章节诊断版报告渲染器。

输入：report_data.json + analysis.json
输出：<公司简称>-运营周报-2026WXX.docx

报告结构：
  封面 + 关键结论卡（3 件最该做的事）
  §1 本周战况速览（KPI 红绿灯）
  §2 流量诊断（漏斗 + 国家归因）
  §3 商品诊断（5-quadrant 标签 + 具体动作）
  §4 关键词诊断（4-quadrant 标签 + 出价建议）
  §5 风险预警（P0/P1/P2 处理项）
  §6 行动 Backlog（P0/P1/P2/P3 任务清单）
"""
from __future__ import annotations

import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


# ----------------------------------------------------------------------
# 颜色 / 样式常量
# ----------------------------------------------------------------------
BRAND_PRIMARY = RGBColor(0x14, 0x55, 0xC0)     # 国际站蓝
BRAND_DARK    = RGBColor(0x0A, 0x2B, 0x60)
COLOR_RED     = RGBColor(0xC0, 0x39, 0x2B)
COLOR_ORANGE  = RGBColor(0xE6, 0x7E, 0x22)
COLOR_GREEN   = RGBColor(0x27, 0xAE, 0x60)
COLOR_GRAY    = RGBColor(0x99, 0x99, 0x99)
COLOR_BG_BLUE = "DCEAFB"
COLOR_BG_GREEN = "DDF1E0"
COLOR_BG_ORANGE = "FCE7CD"
COLOR_BG_RED   = "F8D5D0"
COLOR_BG_GRAY  = "EEEEEE"

PRIO_COLORS = {
    "P0": COLOR_BG_RED,
    "P1": COLOR_BG_ORANGE,
    "P2": COLOR_BG_BLUE,
    "P3": COLOR_BG_GRAY,
}


# ----------------------------------------------------------------------
# 基础工具
# ----------------------------------------------------------------------
def add_p(doc, text="", bold=False, size=10.5, color=None, align=None, italic=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Microsoft YaHei"
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading("", level=level)
    run = h.add_run(text)
    run.font.name = "Microsoft YaHei"
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = BRAND_DARK
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = BRAND_PRIMARY
    else:
        run.font.size = Pt(12)
    return h


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=10, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(display_value(text))
    run.font.size = Pt(size)
    run.font.name = "Microsoft YaHei"
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def display_value(value):
    """报告展示值：接口未返回时明确写“未返回”，不伪装成 0。"""
    if value in (None, ""):
        return "未返回"
    return str(value)


def period_words(meta):
    """根据报告模式返回周期文案，避免月报里出现“本周”。"""
    if (meta or {}).get("mode") == "monthly":
        return {
            "short": "本月",
            "type": "月度",
            "deadline": "本月内",
            "review_hint": "建议每月初阅读、按 backlog 推进。",
            "action_title": "🎯 本月 3 件最该做的事",
            "one_liner": "📊 本月一句话",
            "section1": "一、本月战况速览",
            "backlog_desc": "本月可执行的具体任务清单，按优先级排序。",
            "backlog_total": "本月共",
            "empty_backlog": "🎉 本月 Backlog 为空",
        }
    return {
        "short": "本周",
        "type": "周度",
        "deadline": "本周内",
        "review_hint": "建议每周一阅读、按 backlog 推进。",
        "action_title": "🎯 本周 3 件最该做的事",
        "one_liner": "📊 本周一句话",
        "section1": "一、本周战况速览",
        "backlog_desc": "本周可执行的具体任务清单，按优先级排序。",
        "backlog_total": "本周共",
        "empty_backlog": "🎉 本周 Backlog 为空",
    }


def period_range_text(meta):
    """把 meta 里的开始/结束日期转成报告可读周期。

    Args:
        meta (dict): `prepare_data.py` 写入的报告元信息，通常包含
            `period_start`、`period_end`、`title_period`。

    Returns:
        str: 形如 `2026-04-13 至 2026-04-19` 的周期文案；如果接口或
            命令行没有传入日期，则返回空字符串，避免编造周期。

    Raises:
        None: 该函数只做安全字符串拼接，不主动抛异常。
    """
    meta = meta or {}
    start = meta.get("period_start")
    end = meta.get("period_end")
    if start and end:
        return f"{start} 至 {end}"
    return ""


def make_table(doc, headers, rows, col_widths=None, header_bg=COLOR_BG_BLUE):
    """构造一个表格，自动设置表头底色 + 居中。"""
    if not headers:
        return None
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(hdr_cells[i], header_bg)

    # 数据行
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx in range(n_cols):
            val = row[c_idx] if c_idx < len(row) else ""
            set_cell_text(cells[c_idx], val, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)

    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    return table


def add_callout_box(doc, lines, bg_hex=COLOR_BG_BLUE, accent_color=BRAND_PRIMARY):
    """单格表格做 callout box。lines 为 (text, bold, size, color) 元组列表。"""
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    shade_cell(cell, bg_hex)
    cell.text = ""
    for tup in lines:
        text = tup[0]
        bold = tup[1] if len(tup) > 1 else False
        size = tup[2] if len(tup) > 2 else 10
        color = tup[3] if len(tup) > 3 else None
        p = cell.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = "Microsoft YaHei"
        run.bold = bold
        if color is not None:
            run.font.color.rgb = color
    # 删除空首段
    first_p = cell.paragraphs[0]
    if not first_p.text:
        first_p._element.getparent().remove(first_p._element)


# ----------------------------------------------------------------------
# 封面 + 关键结论卡
# ----------------------------------------------------------------------
def render_cover(doc, meta, analysis, data_quality=None):
    words = period_words(meta)
    # 公司 + 标题
    add_p(doc, " ", size=10)
    add_p(doc, " ", size=10)
    add_p(doc, "阿里巴巴国际站", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, color=COLOR_GRAY)
    add_p(doc, meta.get("company_name", "—"),
          bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER, color=BRAND_DARK)
    title = "运营月报" if meta.get("mode") == "monthly" else "运营周报"
    add_p(doc, title, bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER, color=BRAND_PRIMARY)
    add_p(doc, meta.get("title_period", ""), size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=COLOR_GRAY)
    period_range = period_range_text(meta)
    if period_range:
        add_p(doc, f"报告周期：{period_range}", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=COLOR_GRAY)

    add_p(doc, " ")
    add_p(doc, "—— 数据驱动的运营诊断与行动指南 ——", italic=True, size=10,
          align=WD_ALIGN_PARAGRAPH.CENTER, color=COLOR_GRAY)
    add_p(doc, " ")

    # 一句话战况
    one_liner = analysis.get("one_liner", "")
    add_callout_box(doc, [
        (words["one_liner"], True, 12, BRAND_DARK),
        (one_liner, False, 11, None),
    ], bg_hex=COLOR_BG_BLUE)

    add_p(doc, " ")

    # 数据可信度卡：把接口覆盖情况放在首页，避免读者误把缺失字段当 0。
    dq = data_quality or analysis.get("data_quality") or {}
    checks = dq.get("checks") or {}
    if dq:
        check_text = " ｜ ".join(
            f"{name}:{'已返回' if ok else '未返回'}"
            for name, ok in [
                ("总览", checks.get("summary_indicators")),
                ("漏斗", checks.get("funnel")),
                ("国家", checks.get("region")),
                ("广告", checks.get("ads")),
                ("商品", checks.get("products")),
                ("风险", checks.get("risk")),
            ]
        )
        add_callout_box(doc, [
            ("🧾 数据可信度", True, 12, BRAND_DARK),
            (f"覆盖率 {dq.get('coverage_rate', 0) * 100:.0f}% ｜ 状态：{dq.get('status', '未返回')}", False, 10, None),
            (check_text or "未返回", False, 9, None),
        ], bg_hex=COLOR_BG_BLUE if dq.get("status") == "ok" else COLOR_BG_ORANGE)

    add_p(doc, " ")

    # 3 件最该做的事
    add_callout_box(doc, [
        (words["action_title"], True, 13, BRAND_DARK),
    ], bg_hex=COLOR_BG_GREEN)

    top3 = analysis.get("top3_actions", [])
    if not top3:
        add_p(doc, "  （本期无紧急行动项；详见 §6 行动 Backlog 中的常规任务）",
              italic=True, color=COLOR_GRAY)
    else:
        for i, act in enumerate(top3, 1):
            prio = act.get("priority", "P1")
            color = COLOR_RED if prio == "P0" else (COLOR_ORANGE if prio == "P1" else BRAND_PRIMARY)
            add_p(doc, f"  {i}. 【{prio}｜{act.get('deadline', words['deadline'])}】 {act.get('title', '')}",
                  bold=True, size=11, color=color)
            add_p(doc, f"      └ 原因：{act.get('why', '')}", size=10, color=COLOR_GRAY)
            add_p(doc, f"      └ 操作位置：{act.get('where', '')}", size=10, color=COLOR_GRAY)
            add_p(doc, " ", size=4)

    # 报告说明
    add_p(doc, " ")
    add_p(doc, f"本报告基于阿里巴巴国际站官方接口数据，结合通用电商诊断规则生成，{words['review_hint']}",
          italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=COLOR_GRAY)

    page_break(doc)


# ----------------------------------------------------------------------
# §1 战况速览
# ----------------------------------------------------------------------
def render_section1_kpi(doc, analysis, meta=None):
    add_heading(doc, "一、经营总览", level=1)
    period_range = period_range_text(meta)
    if period_range:
        add_p(doc, f"报告周期：{period_range}", bold=True, color=BRAND_DARK, size=10)
    add_p(doc, "老板视角双基准：同行均值是达标线，同行优秀是增长目标。",
          italic=True, color=COLOR_GRAY, size=10)
    add_p(doc, " ")

    lights = analysis.get("kpi_traffic_lights", [])
    if not lights:
        add_p(doc, "（核心 KPI 数据缺失，请检查 MCP 接口返回）", italic=True, color=COLOR_GRAY)
        page_break(doc)
        return

    headers = ["指标", "本期", "环比", "行业均值", "行业优秀", "诊断"]
    rows = []
    for l in lights:
        crc = l.get("crc", 0)
        crc_str = f"{crc * 100:+.1f}%" if crc else "—"
        rows.append([
            l.get("name", ""),
            display_value(l.get("value")),
            crc_str,
            display_value(l.get("rival_avg")),
            display_value(l.get("rival_good")),
            f"{l.get('light', '')} {l.get('diag', '')}",
        ])
    make_table(doc, headers, rows, col_widths=[2.3, 1.8, 1.6, 1.8, 1.8, 6.5])

    add_p(doc, " ")
    # 一行总结：几个 🟢 / 🟡 / 🔴
    g = sum(1 for l in lights if l.get("light") == "🟢")
    y = sum(1 for l in lights if l.get("light") == "🟡")
    r = sum(1 for l in lights if l.get("light") == "🔴")
    add_callout_box(doc, [
        ("📌 速览结论", True, 11, BRAND_DARK),
        (f"  🟢 表现优异 {g} 项 ｜ 🟡 处于平均 {y} 项 ｜ 🔴 需重点提升 {r} 项",
         False, 10, None),
    ], bg_hex=COLOR_BG_BLUE)

    page_break(doc)


# ----------------------------------------------------------------------
# §2 流量诊断
# ----------------------------------------------------------------------
def render_section2_traffic(doc, analysis):
    add_heading(doc, "三、流量结构与漏斗", level=1)
    add_p(doc, "回答 3 个问题：流量从哪来？为什么涨/跌？漏在哪一段？",
          italic=True, color=COLOR_GRAY, size=10)
    add_p(doc, " ")

    funnel = analysis.get("funnel_diagnosis", {})
    geo = analysis.get("country_channel", {})

    # 2.1 漏斗
    add_heading(doc, "2.1 转化漏斗诊断", level=2)
    totals = funnel.get("totals", {})
    if totals:
        def fmt_count(v):
            return f"{v:,}" if isinstance(v, (int, float)) else "未返回"
        add_p(doc, f"漏斗概况：曝光 {fmt_count(totals.get('imps'))} → 访客 {fmt_count(totals.get('visitor'))} "
                   f"→ 商机 {fmt_count(totals.get('inquiry'))} → 订单 {fmt_count(totals.get('order'))}",
              size=10)
    add_p(doc, " ")

    stages = funnel.get("stages", [])
    if stages:
        headers = ["漏斗段", "进入", "产出", "转化率", "行业均值", "差距"]
        rows = []
        for s in stages:
            rate = s.get("rate", 0)
            base = s.get("baseline", 0)
            diff = (rate - base) / base * 100 if base else 0
            diff_str = f"{diff:+.1f}%" if base else "—"
            rows.append([
                s.get("stage", ""),
                f"{s.get('in', 0):,}",
                f"{s.get('out', 0):,}",
                f"{rate * 100:.2f}%",
                f"{base * 100:.2f}%",
                diff_str,
            ])
        make_table(doc, headers, rows, col_widths=[3.5, 2, 2, 2.2, 2.2, 2])

    add_p(doc, " ")
    anomalies = funnel.get("anomalies", [])
    if anomalies:
        add_callout_box(doc, [
            ("🔍 漏斗反常段诊断", True, 11, BRAND_DARK),
        ] + [(f"  {a['level']} {a['advice']}", False, 10, None) for a in anomalies],
                       bg_hex=COLOR_BG_ORANGE if any(a["level"].startswith("🔴") for a in anomalies) else COLOR_BG_BLUE)
    else:
        add_callout_box(doc, [
            ("✅ 漏斗整体健康", True, 11, COLOR_GREEN),
            ("各段转化率均接近行业均值，无明显反常段。", False, 10, None),
        ], bg_hex=COLOR_BG_GREEN)

    add_p(doc, " ")

    # 2.2 国家归因
    add_heading(doc, "2.2 国家流量归因（涨幅/跌幅 Top3）", level=2)
    risers = geo.get("risers", [])
    fallers = geo.get("fallers", [])

    if risers:
        add_p(doc, "🟢 流量涨幅 Top3", bold=True, size=11, color=COLOR_GREEN)
        rows = [[r["country"], f"{r['uv']:,}", r["crc_str"]] for r in risers]
        make_table(doc, ["国家", "UV", "环比"], rows, col_widths=[5, 3, 3])
        add_p(doc, " ")
    else:
        add_p(doc, "（本期无显著流量涨幅国家）", italic=True, color=COLOR_GRAY, size=10)

    if fallers:
        add_p(doc, "🔴 流量跌幅 Top3", bold=True, size=11, color=COLOR_RED)
        rows = [[f["country"], f"{f['uv']:,}", f["crc_str"]] for f in fallers]
        make_table(doc, ["国家", "UV", "环比"], rows, col_widths=[5, 3, 3])
        add_p(doc, " ")
    else:
        add_p(doc, "（本期无显著流量跌幅国家）", italic=True, color=COLOR_GRAY, size=10)

    insights = geo.get("insights", [])
    actions = geo.get("actions", [])
    add_callout_box(doc, [
        ("💡 国家归因结论", True, 11, BRAND_DARK),
    ] + [(f"  {ins}", False, 10, None) for ins in insights[:5]],
                   bg_hex=COLOR_BG_BLUE)

    add_p(doc, " ")
    if actions:
        add_callout_box(doc, [
            ("🎯 建议动作", True, 11, BRAND_DARK),
        ] + [(f"  • {a}", False, 10, None) for a in actions[:5]],
                       bg_hex=COLOR_BG_GREEN)

    add_p(doc, " ")

    # 2.3 渠道
    add_heading(doc, "2.3 渠道流量分布", level=2)
    channels = geo.get("channels", [])
    if channels:
        rows = [[c["channel"], f"{c['uv']:,}", f"{c['crc'] * 100:+.1f}%", c["vs_rival"]] for c in channels]
        make_table(doc, ["渠道", "UV", "环比", "本店/行业"], rows, col_widths=[5, 3, 2.5, 3])
        alerts = geo.get("channel_alerts", [])
        if alerts:
            add_p(doc, " ")
            add_callout_box(doc, [
                ("🚦 渠道稳定性预警", True, 11, BRAND_DARK),
            ] + [(f"  {a.get('summary')}", False, 9, None) for a in alerts[:5]],
                           bg_hex=COLOR_BG_ORANGE if any(a.get("priority") == "P1" for a in alerts) else COLOR_BG_BLUE)
    else:
        add_p(doc, "（渠道数据不可用）", italic=True, color=COLOR_GRAY)

    page_break(doc)


def render_section_star(doc, report_data):
    """老板版 P0 星级/保星诊断。"""
    add_heading(doc, "二、P0 星级 / 保星诊断", level=1)
    diagnosis = (report_data or {}).get("diagnosis") or {}
    conclusion = diagnosis.get("conclusion") or "星级诊断数据未返回"
    abilities = diagnosis.get("abilities") or []
    advices = diagnosis.get("advices") or []

    bg = COLOR_BG_RED if "降星" in conclusion else COLOR_BG_BLUE
    add_callout_box(doc, [
        ("⭐ 老板先看", True, 12, BRAND_DARK),
        (conclusion, False, 10, None),
    ], bg_hex=bg)
    add_p(doc, " ")

    if abilities:
        rows = []
        for ability in abilities:
            kpis = ability.get("kpis") or []
            evidence = "；".join(
                f"{k.get('name')} {display_value(k.get('value'))} / 下一档 {display_value(k.get('next_level_avg'))}"
                for k in kpis[:2]
            )
            rows.append([
                ability.get("ability"),
                display_value(ability.get("score")),
                display_value(ability.get("star")),
                evidence or "未返回",
            ])
        make_table(doc, ["能力维度", "当前分", "当前星级", "关键缺口"], rows,
                   col_widths=[2.2, 1.6, 1.6, 9.0], header_bg=COLOR_BG_ORANGE)
    else:
        add_p(doc, "（星级能力项未返回）", italic=True, color=COLOR_GRAY)

    if advices:
        add_p(doc, " ")
        lines = [("🎯 平台保星建议", True, 11, BRAND_DARK)]
        for item in advices[:4]:
            details = "；".join(str(x) for x in (item.get("details") or [])[:2])
            lines.append((f"  • {item.get('indicator')}：{details}", False, 9, None))
        add_callout_box(doc, lines, bg_hex=COLOR_BG_BLUE)

    page_break(doc)


# ----------------------------------------------------------------------
# §3 商品诊断（核心）
# ----------------------------------------------------------------------
def render_quadrant_card(doc, tag_key, tag_meta, items, max_show=10):
    """渲染一个 quadrant 卡片：标签 emoji + 名称 + 列表"""
    if not items:
        add_p(doc, f"  {tag_meta['emoji']} {tag_meta['label']}：本期无", italic=True, color=COLOR_GRAY)
        return

    bg = {
        "green": COLOR_BG_GREEN, "orange": COLOR_BG_ORANGE,
        "red": COLOR_BG_RED, "gray": COLOR_BG_GRAY, "blue": COLOR_BG_BLUE,
    }.get(tag_meta["color"], COLOR_BG_BLUE)
    color = {
        "green": COLOR_GREEN, "orange": COLOR_ORANGE,
        "red": COLOR_RED, "gray": COLOR_GRAY, "blue": BRAND_PRIMARY,
    }.get(tag_meta["color"], BRAND_PRIMARY)

    add_callout_box(doc, [
        (f"{tag_meta['emoji']} {tag_meta['label']} ({len(items)} 款)", True, 12, color),
    ], bg_hex=bg)

    for i, p in enumerate(items[:max_show], 1):
        product_id = display_value(p.get("product_id"))
        add_p(doc, f"  {i}. [商品ID: {product_id}] {p.get('title', '')}", bold=True, size=10)
        if p.get("imps") or p.get("fb_num"):
            line = f"     曝光 {p.get('imps', 0):,} ｜ 询盘 {p.get('fb_num', 0)} ｜ 询盘率 {p.get('fb_rate_str', '—')}"
            add_p(doc, line, size=9, color=COLOR_GRAY)
        if p.get("why"):
            add_p(doc, f"     ▶ 诊断：{p['why']}", size=9, color=BRAND_DARK)
        if p.get("actions"):
            for a in p["actions"]:
                add_p(doc, f"        • {a}", size=9)
        add_p(doc, " ", size=4)

    if len(items) > max_show:
        add_p(doc, f"  …… 另有 {len(items) - max_show} 款同类商品（数据见报告附录）",
              italic=True, color=COLOR_GRAY, size=9)
    add_p(doc, " ")


def render_section3_products(doc, analysis, report_data=None):
    add_heading(doc, "四、商品结构与 Top 商品处理清单", level=1)
    report_data = report_data or {}
    is_monthly = ((report_data.get("meta") or {}).get("mode") == "monthly")
    intro = "月报模式展示商品供给、橱窗和类目结构；若接口返回商品级表现，再补充分群诊断。" if is_monthly else "把 Top 商品按「曝光-转化」双维度自动分到 5 个 quadrant，每款给出具体动作。"
    add_p(doc, intro,
          italic=True, color=COLOR_GRAY, size=10)
    add_p(doc, " ")

    quadrants = analysis.get("products_quadrant", {})
    meta_map = analysis.get("products_quadrant_meta", {})
    products_data = report_data.get("products") or {}
    market = report_data.get("market") or {}
    product_selection = market.get("product_selection_recent_30d") or []
    overview = products_data.get("overview") or {}
    product_layers = overview.get("product_layers") or []
    top5_categories = products_data.get("top5_categories") or []

    if product_layers:
        total_products = overview.get("total_products")
        add_callout_box(doc, [
            ("📦 商品分层诊断", True, 11, BRAND_DARK),
            (f"  全店商品数 {display_value(total_products)}；老板重点看普通品是否过剩、爆品是否太少、潜力品是否能孵化。",
             False, 10, None),
        ], bg_hex=COLOR_BG_BLUE)
        rows = []
        for layer in product_layers:
            rows.append([
                layer.get("stage_name"),
                display_value(layer.get("prod_cnt")),
                f"{float(layer.get('prod_cnt_ratio')) * 100:.1f}%" if layer.get("prod_cnt_ratio") not in (None, "") else "未返回",
                f"{float(layer.get('cate_avg_ratio')) * 100:.1f}%" if layer.get("cate_avg_ratio") not in (None, "") else "未返回",
                display_value(layer.get("avg_uv_30d")),
                display_value(layer.get("avg_inquiry_90d")),
            ])
        make_table(doc, ["层级", "商品数", "占比", "同品类均值", "30天均UV", "90天均询盘"], rows,
                   col_widths=[2.2, 1.6, 1.6, 2.0, 1.8, 2.0], header_bg=COLOR_BG_BLUE)
        add_p(doc, " ")

    if top5_categories:
        add_heading(doc, "4.1 Top5 类目询盘效率", level=2)
        rows = []
        for c in top5_categories[:5]:
            rate = c.get("inquiries_rate")
            rate_str = f"{float(rate) * 100:.2f}%" if rate not in (None, "") else "未返回"
            rows.append([
                c.get("type"),
                display_value(c.get("visitors")),
                display_value(c.get("inquiries")),
                rate_str,
            ])
        make_table(doc, ["类目", "UV", "询盘", "询盘率"], rows,
                   col_widths=[7.5, 1.8, 1.8, 2.0], header_bg=COLOR_BG_GREEN)
        add_p(doc, " ")

    if is_monthly:
        categories = products_data.get("categories") or []
        shelf_products = products_data.get("shelf_products") or []
        add_callout_box(doc, [
            ("📦 月度商品供给概览", True, 11, BRAND_DARK),
            (f"  橱窗商品 {products_data.get('shelf_total', '未返回')} 款 ｜ 类目结构 {len(categories)} 个 ｜ 商品级曝光/商机接口：{'已返回' if any(quadrants.values()) else '未返回'}",
             False, 10, None),
        ], bg_hex=COLOR_BG_BLUE)
        add_p(doc, " ")
        if categories:
            make_table(
                doc,
                ["一级类目", "二级类目", "三级类目", "商品数"],
                [[c.get("lv1"), c.get("lv2"), c.get("lv3"), c.get("count")] for c in categories[:8]],
                col_widths=[3.2, 3.2, 4.2, 1.6],
            )
            add_p(doc, " ")
        if shelf_products:
            make_table(
                doc,
                ["橱窗位", "商品ID", "商品名称", "类目"],
                [[p.get("position"), display_value(p.get("product_id")), p.get("name"), p.get("category")] for p in shelf_products[:10]],
                col_widths=[1.4, 2.4, 7.0, 3.2],
            )
            add_p(doc, " ")
        if not any(quadrants.values()):
            add_callout_box(doc, [
                ("说明", True, 10, BRAND_DARK),
                ("月报接口本期未返回商品级曝光/商机 Top 表，因此不强行生成印钞款/失血款结论；请用周报查看商品四象限。", False, 9, None),
            ], bg_hex=COLOR_BG_ORANGE)
        if product_selection:
            add_p(doc, " ")
            add_heading(doc, "3.1 近 30 天行业选品机会", level=2)
            add_p(doc, "以下为行业滚动近 30 天口径，用于选品/补品参考，不等同于本月店铺经营 KPI。",
                  italic=True, color=COLOR_GRAY, size=9)
            rows = [
                [
                    display_value(p.get("product_id")),
                    p.get("product_name"),
                    p.get("cate_name"),
                    p.get("price"),
                    p.get("moq"),
                    p.get("ab_cnt_30d"),
                    p.get("order_cnt_30d"),
                    p.get("detail_url") or "未返回",
                ]
                for p in product_selection[:8]
            ]
            make_table(doc, ["商品ID", "行业商品", "类目", "价格", "MOQ", "近30天询盘", "近30天订单", "详情链接"], rows,
                       col_widths=[2.1, 4.0, 2.0, 1.5, 1.1, 1.5, 1.5, 3.2])
            page_break(doc)
            return

    # Top 商品分群概览
    counts = {k: len(v) for k, v in quadrants.items()}
    add_callout_box(doc, [
        ("📊 Top 商品处理概览", True, 11, BRAND_DARK),
        (f"  🔥 印钞款 {counts.get('ink_print', 0)}  ｜  ⚡ 潜力款 {counts.get('potential', 0)}"
         f"  ｜  🩹 失血款 {counts.get('bleeding', 0)}"
         f"  ｜  🪦 僵尸款 {counts.get('zombie', 0)}  ｜  ❓ 观察款 {counts.get('watch', 0)}",
         False, 10, None),
    ], bg_hex=COLOR_BG_BLUE)
    add_p(doc, " ")

    order = ["ink_print", "potential", "bleeding", "zombie", "watch"]
    for tag in order:
        meta = meta_map.get(tag, {"emoji": "•", "label": tag, "color": "blue"})
        render_quadrant_card(doc, tag, meta, quadrants.get(tag, []), max_show=5)

    if product_selection:
        add_heading(doc, "3.1 近 30 天行业选品机会", level=2)
        add_p(doc, "行业滚动近 30 天口径，用于补品和主推方向参考。",
              italic=True, color=COLOR_GRAY, size=9)
        rows = [
            [
                display_value(p.get("product_id")),
                p.get("product_name"),
                p.get("price"),
                p.get("moq"),
                p.get("ab_cnt_30d"),
                p.get("order_cnt_30d"),
                p.get("detail_url") or "未返回",
            ]
            for p in product_selection[:8]
        ]
        make_table(doc, ["商品ID", "行业商品", "价格", "MOQ", "询盘", "订单", "详情链接"], rows,
                   col_widths=[2.1, 4.5, 1.5, 1.1, 1.3, 1.3, 3.4])

    page_break(doc)


# ----------------------------------------------------------------------
# §4 关键词诊断
# ----------------------------------------------------------------------
def render_keyword_quadrant(doc, tag_key, meta, items, max_show=10):
    if not items:
        add_p(doc, f"  {meta['emoji']} {meta['label']}：本期无", italic=True, color=COLOR_GRAY)
        return

    bg = {
        "green": COLOR_BG_GREEN, "orange": COLOR_BG_ORANGE,
        "red": COLOR_BG_RED, "blue": COLOR_BG_BLUE,
    }.get(meta["color"], COLOR_BG_BLUE)
    color = {
        "green": COLOR_GREEN, "orange": COLOR_ORANGE,
        "red": COLOR_RED, "blue": BRAND_PRIMARY,
    }.get(meta["color"], BRAND_PRIMARY)

    add_callout_box(doc, [
        (f"{meta['emoji']} {meta['label']} ({len(items)} 个)", True, 12, color),
    ], bg_hex=bg)

    headers = ["关键词", "来源", "排名/证据", "询盘", "花费($)", "对应商品"]
    if tag_key == "expand":
        headers = ["关键词", "来源", "排名/证据"]

    rows = []
    for k in items[:max_show]:
        if tag_key == "expand":
            rows.append([
                k.get("keyword", ""),
                k.get("source_tag") or "行业/店铺词表",
                k.get("rank_label") or (f"行业曝光 {k.get('industry_imps', 0):,}" if k.get("industry_imps") else "榜单词"),
            ])
        else:
            rows.append([
                k.get("keyword", ""),
                k.get("source_tag") or "P4P/店铺词表",
                k.get("rank_label") or (f"排名 {k.get('rank')}" if k.get("rank") else "榜单词"),
                f"{k.get('inquiry', 0)}",
                f"{k.get('cost', 0):.2f}" if k.get("cost") else "—",
                k.get("product_name") or "—",
            ])
    make_table(doc, headers, rows)
    add_p(doc, " ")

    # 每个词的诊断和动作
    for k in items[:max_show]:
        add_p(doc, f"  ▸ 「{k.get('keyword', '')}」", bold=True, size=10, color=color)
        if k.get("why"):
            add_p(doc, f"     诊断：{k['why']}", size=9, color=COLOR_GRAY)
        if k.get("actions"):
            for a in k["actions"]:
                add_p(doc, f"     • {a}", size=9)
        add_p(doc, " ", size=4)
    add_p(doc, " ")


def render_section4_keywords(doc, analysis, report_data=None):
    report_data = report_data or {}
    add_heading(doc, "附录 A：关键词与投流机会", level=1)
    add_p(doc, "把 P4P 关键词与自然词按「产出-成本」分到 4 个 quadrant，每个词给出具体出价/动作。",
          italic=True, color=COLOR_GRAY, size=10)
    add_p(doc, " ")

    quadrants = analysis.get("keywords_quadrant", {})
    meta_map = analysis.get("keywords_quadrant_meta", {})

    # 概览
    counts = {k: len(v) for k, v in quadrants.items()}
    add_callout_box(doc, [
        ("📊 关键词分群概览", True, 11, BRAND_DARK),
        (f"  ⭐ 金主词 {counts.get('gold', 0)}  ｜  💰 烧钱词 {counts.get('burning', 0)}"
         f"  ｜  🚀 潜力词 {counts.get('potential', 0)}  ｜  🌱 拓展词 {counts.get('expand', 0)}",
         False, 10, None),
    ], bg_hex=COLOR_BG_BLUE)
    add_p(doc, " ")

    order = ["gold", "burning", "potential", "expand"]
    for tag in order:
        meta = meta_map.get(tag, {"emoji": "•", "label": tag, "color": "blue"})
        render_keyword_quadrant(doc, tag, meta, quadrants.get(tag, []), max_show=5)

    market = report_data.get("market") or {}
    keyword_market = market.get("keyword_market") or []
    next_auction = market.get("next_month_auction") or []
    behaviors = market.get("behavior_semantics") or []

    if keyword_market:
        add_heading(doc, "4.1 行业热词与品牌广告词机会", level=2)
        rows = [
            [
                k.get("keyword"),
                k.get("biz_line"),
                k.get("channel"),
                k.get("year_imps_index"),
                k.get("business_rate"),
                k.get("sell_status"),
                k.get("tags"),
            ]
            for k in keyword_market[:10]
        ]
        make_table(doc, ["关键词", "业务线", "渠道", "曝光指数", "商机转化率", "售卖状态", "标签"], rows,
                   col_widths=[3.0, 1.4, 1.2, 1.4, 1.5, 2.5, 4.0])
        add_p(doc, " ")

    if next_auction:
        add_heading(doc, "4.2 次月可打标/竞价资源", level=2)
        add_p(doc, "这是次月资源口径，用于提前锁词，不代表本报告周期表现。",
              italic=True, color=COLOR_GRAY, size=9)
        rows = [
            [
                k.get("keyword"),
                k.get("biz_line"),
                k.get("year_imps_index"),
                k.get("business_rate"),
                k.get("sell_status"),
                k.get("price"),
            ]
            for k in next_auction[:8]
        ]
        make_table(doc, ["关键词", "业务线", "曝光指数", "商机转化率", "售卖状态", "价格"], rows,
                   col_widths=[3.2, 1.4, 1.5, 1.8, 4.0, 2.0])
        add_p(doc, " ")

    if behaviors:
        add_callout_box(doc, [
            ("🧭 站内行为背景", True, 11, BRAND_DARK),
        ] + [(f"  {x}", False, 9, None) for x in behaviors[:5]], bg_hex=COLOR_BG_BLUE)

    page_break(doc)


# ----------------------------------------------------------------------
# §5 风险预警
# ----------------------------------------------------------------------
def render_section5_risk(doc, analysis):
    add_heading(doc, "附录 B：合规风险", level=1)
    add_p(doc, "店铺合规与处罚风险，按优先级处理。",
          italic=True, color=COLOR_GRAY, size=10)
    add_p(doc, " ")

    risks = analysis.get("risks", [])
    if not risks:
        add_callout_box(doc, [
            ("✅ 本期无风险预警", True, 12, COLOR_GREEN),
            ("店铺显性处罚、知产、订单风险未触发；下表列出本次接口返回的健康项。", False, 10, None),
        ], bg_hex=COLOR_BG_GREEN)
        risk_health = analysis.get("risk_health", [])
        if risk_health:
            add_p(doc, " ")
            headers = ["检查项", "返回值", "状态", "判断标准", "入口"]
            rows = [
                [
                    r.get("name", ""),
                    display_value(r.get("value")),
                    r.get("status", ""),
                    r.get("standard", ""),
                    r.get("where", ""),
                ]
                for r in risk_health
            ]
            make_table(doc, headers, rows, col_widths=[2.3, 1.8, 1.8, 5.2, 4.3], header_bg=COLOR_BG_GREEN)
        page_break(doc)
        return

    headers = ["优先级", "风险项", "数量", "推荐动作", "处理位置"]
    rows = []
    for r in risks:
        rows.append([r["priority"], r["name"], str(r["count"]), r["action"], r["where"]])
    make_table(doc, headers, rows, col_widths=[1.5, 2.5, 1.2, 5.5, 4.5])

    page_break(doc)


def render_section_service(doc, report_data):
    """服务力老板摘要。只展示响应纪律，不展开沟通内容。"""
    add_heading(doc, "五、服务力预警", level=1)
    service = (report_data or {}).get("service") or {}
    if not service:
        add_callout_box(doc, [
            ("服务数据未接入", True, 12, COLOR_ORANGE),
            ("本期未返回服务响应摘要；老板版完整周报建议补采 5 分钟回复率、平均回复时长和 12h+ 回复条数。", False, 10, None),
        ], bg_hex=COLOR_BG_ORANGE)
        page_break(doc)
        return

    first_5 = service.get("first_5min_reply_rate_30d")
    first_5_str = f"{first_5 * 100:.2f}%" if isinstance(first_5, (int, float)) else "未返回"
    avg_reply = service.get("avg_reply_time_30d")
    avg_reply_str = f"{avg_reply:.2f} h" if isinstance(avg_reply, (int, float)) else "未返回"
    rows = [
        ["30天首次5分钟回复率", first_5_str, "对照平台/历史/用户目标"],
        ["30天平均回复时长", avg_reply_str, "对照平台/历史/用户目标"],
        ["上周 12h+ 回复条数", display_value(service.get("reply_over_12h_count")), "按买家优先级和平台规则复查"],
        ["离线消息条数", display_value(service.get("offline_msg_count")), "当天处理"],
        ["未跟进/重复回复", f"{display_value(service.get('not_follow_count'))} / {display_value(service.get('repeat_reply_count'))}", "周内清零"],
    ]
    make_table(doc, ["指标", "本店", "老板判断线"], rows, col_widths=[4.0, 3.0, 6.0],
               header_bg=COLOR_BG_RED if service.get("warnings") else COLOR_BG_GREEN)

    warnings = service.get("warnings") or []
    add_p(doc, " ")
    if warnings:
        add_callout_box(doc, [
            ("🔴 服务响应动作", True, 12, COLOR_RED),
        ] + [(f"  • {w}", False, 10, None) for w in warnings[:5]], bg_hex=COLOR_BG_RED)
    else:
        add_callout_box(doc, [
            ("✅ 服务力未触发红色预警", True, 12, COLOR_GREEN),
            ("继续保持响应纪律，并把优秀响应节奏复制给所有接待班次。", False, 10, None),
        ], bg_hex=COLOR_BG_GREEN)

    page_break(doc)


# ----------------------------------------------------------------------
# §6 行动 Backlog
# ----------------------------------------------------------------------
def render_section6_backlog(doc, analysis, meta=None):
    words = period_words(meta)
    add_heading(doc, "六、行动 Backlog", level=1)
    add_p(doc, words["backlog_desc"],
          italic=True, color=COLOR_GRAY, size=10)
    add_p(doc, " ")

    backlog = analysis.get("backlog", {})

    prio_titles = {
        "P0": "🔴 P0 紧急（今天做）",
        "P1": f"🟠 P1 高优（{words['deadline']}做完）",
        "P2": "🔵 P2 中优（下个周期做）",
        "P3": "⚪ P3 低优（本月做）",
    }
    prio_colors = {
        "P0": COLOR_RED, "P1": COLOR_ORANGE, "P2": BRAND_PRIMARY, "P3": COLOR_GRAY,
    }
    prio_bgs = {
        "P0": COLOR_BG_RED, "P1": COLOR_BG_ORANGE, "P2": COLOR_BG_BLUE, "P3": COLOR_BG_GRAY,
    }

    total = sum(len(backlog.get(p, [])) for p in ["P0", "P1", "P2", "P3"])
    add_callout_box(doc, [
        ("📋 Backlog 概览", True, 12, BRAND_DARK),
        (f"  {words['backlog_total']} {total} 项任务："
         f"P0 {len(backlog.get('P0', []))} 项 ｜ P1 {len(backlog.get('P1', []))} 项"
         f" ｜ P2 {len(backlog.get('P2', []))} 项 ｜ P3 {len(backlog.get('P3', []))} 项",
         False, 10, None),
    ], bg_hex=COLOR_BG_BLUE)
    add_p(doc, " ")

    for prio in ["P0", "P1", "P2", "P3"]:
        items = backlog.get(prio, [])
        if not items:
            continue
        add_callout_box(doc, [(prio_titles[prio], True, 12, prio_colors[prio])],
                       bg_hex=prio_bgs[prio])
        for i, item in enumerate(items, 1):
            add_p(doc, f"  ☐ {i}. {item}", size=10)
        add_p(doc, " ")

    if total == 0:
        add_callout_box(doc, [
            (words["empty_backlog"], True, 12, COLOR_GREEN),
            ("店铺各项指标平稳，无紧急行动项。建议持续关注 §3 商品 / §4 关键词的优化空间。",
             False, 10, None),
        ], bg_hex=COLOR_BG_GREEN)


# ----------------------------------------------------------------------
# 主入口
# ----------------------------------------------------------------------
def build(report_data: dict, analysis: dict, output_path: Path) -> Path:
    doc = Document()

    # 默认样式
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)

    meta = report_data.get("meta", {}) or {}

    render_cover(doc, meta, analysis, report_data.get("data_quality"))
    render_section1_kpi(doc, analysis, meta)
    render_section_star(doc, report_data)
    render_section2_traffic(doc, analysis)
    render_section3_products(doc, analysis, report_data)
    render_section_service(doc, report_data)
    render_section6_backlog(doc, analysis, meta)
    render_section4_keywords(doc, analysis, report_data)
    render_section5_risk(doc, analysis)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def main():
    if len(sys.argv) < 3:
        print("usage: build_docx.py <report_data.json> <analysis.json> [<output.docx>]", file=sys.stderr)
        sys.exit(1)
    rd = json.loads(Path(sys.argv[1]).read_text(encoding="utf-8"))
    an = json.loads(Path(sys.argv[2]).read_text(encoding="utf-8"))
    if len(sys.argv) > 3:
        out = Path(sys.argv[3])
    else:
        meta = rd.get("meta", {})
        company = meta.get("company_name", "店铺")
        period = meta.get("title_period", datetime.now().strftime("%Y-%m-%d"))
        title = "运营月报" if meta.get("mode") == "monthly" else "运营周报"
        out = Path(sys.argv[1]).parent / f"{company}-{title}-{period}.docx"
    build(rd, an, out)
    print(f"docx written: {out}")
    print(f"  size: {out.stat().st_size} bytes")
    d = Document(str(out))
    print(f"  paragraphs: {len(d.paragraphs)}, tables: {len(d.tables)}")


if __name__ == "__main__":
    main()
